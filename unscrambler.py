"""
Copyright 2026 Muhammad Bima Ramadhan

Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation
files (the “Software”), to deal in the Software without restriction, including without limitation the rights to use, copy,
modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons to whom the Software
is furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES
OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE
LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR
IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.
"""

import argparse
import csv
import json
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from typing import TypedDict, cast

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Length, Pt, RGBColor

from scrambler import KeyEntry

TREATMENT_META = {
    "T1": {"rag_status": "RAG-", "refine_status": "Refine-"},
    "T2": {"rag_status": "RAG+", "refine_status": "Refine-"},
    "T3": {"rag_status": "RAG-", "refine_status": "Refine+"},
    "T4": {"rag_status": "RAG+", "refine_status": "Refine+"},
}


@dataclass
class _CLIArgs:
    evaluated_csv: str
    key_json: str
    csv_out: str
    docx_out: str


class LongRow(TypedDict):
    idiom_id: int
    source_text: str
    target_idiom: str
    treatment: str
    rag_status: str
    refine_status: str
    translation: str
    accuracy: float
    acceptability: float
    readability: float
    note: str


def create_styled_cell(
    cell,  # pyright: ignore[reportUnknownParameterType, reportMissingParameterType]
    text: object,
    width: Length,
    bg_hex: str | None = None,
    bold: bool = False,
    font_size: int = 9,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    text_color: RGBColor | None = None,
) -> None:
    cell.width = width
    if bg_hex:
        tcPr = cell._element.get_or_add_tcPr()  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}"/>'))  # pyright: ignore[reportUnknownMemberType]

    p = cell.paragraphs[0]  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)  # pyright: ignore[reportUnknownMemberType]
    p.paragraph_format.space_after = Pt(2)  # pyright: ignore[reportUnknownMemberType]

    run = p.add_run(str(text))  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
    run.font.name = "Times New Roman"  # pyright: ignore[reportUnknownMemberType]
    run.font.size = Pt(font_size)  # pyright: ignore[reportUnknownMemberType]
    run.bold = bold
    if text_color:
        run.font.color.rgb = text_color  # pyright: ignore[reportUnknownMemberType]


def add_excerpt_with_all_bold_idioms(
    cell,  # pyright: ignore[reportUnknownParameterType, reportMissingParameterType]
    source_text: str,
    target_idioms: list[str],
    width: Length,
    font_size: int = 9,
) -> None:
    cell.width = width
    p = cell.paragraphs[0]  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
    p.paragraph_format.space_before = Pt(2)  # pyright: ignore[reportUnknownMemberType]
    p.paragraph_format.space_after = Pt(2)  # pyright: ignore[reportUnknownMemberType]

    valid_idioms = sorted([i for i in target_idioms if i], key=len, reverse=True)
    if not valid_idioms:
        run = p.add_run(source_text)  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        run.font.name = "Times New Roman"  # pyright: ignore[reportUnknownMemberType]
        run.font.size = Pt(font_size)  # pyright: ignore[reportUnknownMemberType]
        return

    pattern = re.compile("|".join(re.escape(i) for i in valid_idioms), re.IGNORECASE)

    last_idx = 0
    for match in pattern.finditer(source_text):
        start, end = match.span()

        if start > last_idx:
            run = p.add_run(source_text[last_idx:start])  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
            run.font.name = "Times New Roman"  # pyright: ignore[reportUnknownMemberType]
            run.font.size = Pt(font_size)  # pyright: ignore[reportUnknownMemberType]

        run_bold = p.add_run(source_text[start:end])  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        run_bold.font.name = "Times New Roman"  # pyright: ignore[reportUnknownMemberType]
        run_bold.font.size = Pt(font_size)  # pyright: ignore[reportUnknownMemberType]
        run_bold.bold = True

        last_idx = end

    if last_idx < len(source_text):
        run = p.add_run(source_text[last_idx:])  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        run.font.name = "Times New Roman"  # pyright: ignore[reportUnknownMemberType]
        run.font.size = Pt(font_size)  # pyright: ignore[reportUnknownMemberType]


def generate_word_tables(evaluated_csv: str, key_json: str, out_file: str):
    with open(key_json, "r", encoding="utf-8") as f:
        key_mapping = cast("dict[str, KeyEntry]", json.load(f))

    grouped_excerpts: defaultdict[str, list[dict[str, str]]] = defaultdict(list)
    with open(evaluated_csv, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            grouped_excerpts[row["source_text"]].append(row)

    doc = docx.Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.0)
    section.page_height = Inches(8.5)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Table 1: Global Translation Matrix")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    t1_table = doc.add_table(rows=1, cols=6)
    t1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1_widths = [
        Inches(0.6),
        Inches(2.2),
        Inches(1.6),
        Inches(1.6),
        Inches(1.6),
        Inches(1.6),
    ]

    header_tr1 = t1_table.rows[0]._element.get_or_add_trPr()  # pyright: ignore[reportPrivateUsage]
    header_tr1.append(parse_xml(f"<w:tblHeader {nsdecls('w')}/>"))  # pyright: ignore[reportUnknownMemberType]

    hdr1_cells = t1_table.rows[0].cells
    hdr1_titles = [
        "No.",
        "Source Text Excerpt",
        "T1 (RAG-, Refine-)",
        "T2 (RAG+, Refine-)",
        "T3 (RAG-, Refine+)",
        "T4 (RAG+, Refine+)",
    ]
    for idx, title in enumerate(hdr1_titles):
        align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        create_styled_cell(
            hdr1_cells[idx],
            title,
            t1_widths[idx],
            bg_hex="1F4E78",
            bold=True,
            font_size=9,
            align=align,
            text_color=RGBColor(0xFF, 0xFF, 0xFF),
        )

    excerpt_num = 1
    excerpt_id_map: dict[str, int] = {}

    for source_text, idiom_entries in grouped_excerpts.items():
        excerpt_id_map[source_text] = excerpt_num
        first_pid = idiom_entries[0]["pair_id"]
        if first_pid not in key_mapping:
            continue
        mapping = key_mapping[first_pid]["mapping"]

        tx_by_code: dict[str, str] = {}
        for col, tx_code in mapping.items():
            tx_by_code[tx_code] = idiom_entries[0].get(f"translation_{col}", "")

        row_cells = t1_table.add_row().cells
        create_styled_cell(
            row_cells[0],
            str(excerpt_num),
            t1_widths[0],
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_excerpt_with_all_bold_idioms(
            row_cells[1],
            source_text,
            [
                entry["target_idiom"]
                for entry in idiom_entries
                if "target_idiom" in entry
            ],
            t1_widths[1],
        )
        create_styled_cell(row_cells[2], tx_by_code.get("T1", ""), t1_widths[2])
        create_styled_cell(row_cells[3], tx_by_code.get("T2", ""), t1_widths[3])
        create_styled_cell(row_cells[4], tx_by_code.get("T3", ""), t1_widths[4])
        create_styled_cell(row_cells[5], tx_by_code.get("T4", ""), t1_widths[5])

        excerpt_num += 1

    _ = doc.add_page_break()

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Table 2: Idiom Grading and Evaluation Sheet")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    t2_table = doc.add_table(rows=1, cols=7)
    t2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2_widths = [
        Inches(1.8),
        Inches(0.8),
        Inches(0.9),
        Inches(0.6),
        Inches(0.6),
        Inches(0.6),
        Inches(3.9),
    ]

    header_tr2 = t2_table.rows[0]._element.get_or_add_trPr()  # pyright: ignore[reportPrivateUsage]
    header_tr2.append(parse_xml(f"<w:tblHeader {nsdecls('w')}/>"))  # pyright: ignore[reportUnknownMemberType]

    hdr2_cells = t2_table.rows[0].cells
    hdr2_titles = [
        "Target Idiom",
        "Excerpt No.",
        "Treatment",
        "Acc.",
        "Accp.",
        "Read.",
        "Note",
    ]
    for idx, title in enumerate(hdr2_titles):
        align = (
            WD_ALIGN_PARAGRAPH.CENTER
            if idx in [1, 2, 3, 4, 5]
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        create_styled_cell(
            hdr2_cells[idx],
            title,
            t2_widths[idx],
            bg_hex="1F4E78",
            bold=True,
            font_size=9,
            align=align,
            text_color=RGBColor(0xFF, 0xFF, 0xFF),
        )

    for source_text, idiom_entries in grouped_excerpts.items():
        e_num = excerpt_id_map.get(source_text, "")

        for entry in idiom_entries:
            target_idiom = entry["target_idiom"]
            pid = entry["pair_id"]
            if pid not in key_mapping:
                continue
            mapping = key_mapping[pid]["mapping"]

            first_row_cells = None
            last_row_cells = None

            for col, tx_code in sorted(mapping.items(), key=lambda x: x[1]):
                row_cells = t2_table.add_row().cells

                acc = int(entry.get(f"accuracy_{col}", 0))
                accp = int(entry.get(f"acceptability_{col}", 0))
                read = int(entry.get(f"readability_{col}", 0))

                note_str = entry.get("note", "")
                if note_str and "{" in note_str:
                    note_str = note_str.format(**mapping)

                if first_row_cells is None:
                    first_row_cells = row_cells
                    create_styled_cell(row_cells[0], target_idiom, t2_widths[0])
                    create_styled_cell(
                        row_cells[1],
                        str(e_num),
                        t2_widths[1],
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                    )
                    create_styled_cell(row_cells[6], note_str, t2_widths[6])
                else:
                    create_styled_cell(row_cells[0], "", t2_widths[0])
                    create_styled_cell(row_cells[1], "", t2_widths[1])
                    create_styled_cell(row_cells[6], "", t2_widths[6])

                create_styled_cell(
                    row_cells[2], tx_code, t2_widths[2], align=WD_ALIGN_PARAGRAPH.CENTER
                )
                create_styled_cell(
                    row_cells[3],
                    str(acc),
                    t2_widths[3],
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                create_styled_cell(
                    row_cells[4],
                    str(accp),
                    t2_widths[4],
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                create_styled_cell(
                    row_cells[5],
                    str(read),
                    t2_widths[5],
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                last_row_cells = row_cells

            if first_row_cells is not None and last_row_cells is not None:
                _ = first_row_cells[0].merge(last_row_cells[0])
                _ = first_row_cells[1].merge(last_row_cells[1])
                _ = first_row_cells[6].merge(last_row_cells[6])

                for col_idx in (0, 1, 6):
                    merged_cell = first_row_cells[col_idx]
                    for p in merged_cell.paragraphs[1:]:
                        merged_cell._element.remove(p._element)  # pyright: ignore[reportUnknownMemberType, reportPrivateUsage]

    doc.save(out_file)
    print(f"Word table generated: {out_file}")


def generate_long_csv(evaluated_csv: str, key_json: str, out_file: str):
    with open(key_json, "r", encoding="utf-8") as f:
        key_mapping = cast("dict[str, KeyEntry]", json.load(f))

    long_rows: list[LongRow] = []

    with open(evaluated_csv, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            pid = row["pair_id"]
            if pid not in key_mapping:
                continue

            mapping = key_mapping[pid]["mapping"]

            try:
                scores = {
                    "A": {
                        "acc": float(row["accuracy_A"]),
                        "accp": float(row["acceptability_A"]),
                        "read": float(row["readability_A"]),
                    },
                    "B": {
                        "acc": float(row["accuracy_B"]),
                        "accp": float(row["acceptability_B"]),
                        "read": float(row["readability_B"]),
                    },
                    "C": {
                        "acc": float(row["accuracy_C"]),
                        "accp": float(row["acceptability_C"]),
                        "read": float(row["readability_C"]),
                    },
                    "D": {
                        "acc": float(row["accuracy_D"]),
                        "accp": float(row["acceptability_D"]),
                        "read": float(row["readability_D"]),
                    },
                }
            except (ValueError, KeyError, TypeError):
                print(
                    f"Warning: Missing or malformed data at pair_id {pid}. Skipping row.",
                    file=sys.stderr,
                )
                continue

            for col, tx in sorted(mapping.items(), key=lambda i: i[1]):
                acc = scores[col]["acc"]
                accp = scores[col]["accp"]
                read = scores[col]["read"]

                long_rows.append(
                    LongRow(
                        idiom_id=int(pid),
                        source_text=row["source_text"],
                        target_idiom=row["target_idiom"],
                        treatment=tx,
                        rag_status=TREATMENT_META[tx]["rag_status"],
                        refine_status=TREATMENT_META[tx]["refine_status"],
                        translation=row[f"translation_{col}"],
                        accuracy=acc,
                        acceptability=accp,
                        readability=read,
                        note=tx == "T1" and row["note"].format(**mapping) or "",
                    )
                )

    long_headers = [
        "idiom_id",
        "source_text",
        "target_idiom",
        "treatment",
        "rag_status",
        "refine_status",
        "translation",
        "accuracy",
        "acceptability",
        "readability",
        "note",
    ]

    with open(out_file, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=long_headers)
        writer.writeheader()
        writer.writerows(long_rows)

    print(f"Data successfully unscrambled and pivoted to long format: {out_file}")


def main():
    parser = argparse.ArgumentParser(
        description="Decode evaluation matrix and pivot to long format for R analysis and Word docx for reporting."
    )
    _ = parser.add_argument("evaluated_csv", help="Path to the scored blind_test.csv")
    _ = parser.add_argument("key_json", help="Path to the blind_key.json mapping key")
    _ = parser.add_argument(
        "-co",
        "--csv-out",
        default="translations_long.csv",
        help="Path to save the unscrambled long-format dataset",
    )
    _ = parser.add_argument(
        "-wo",
        "-do",
        "--word-out",
        "--docx-out",
        default="grading_sheet.docx",
        dest="docx_out",
        help="Path to save the Word table",
    )
    args = parser.parse_args(namespace=_CLIArgs)
    generate_long_csv(args.evaluated_csv, args.key_json, args.csv_out)
    generate_word_tables(args.evaluated_csv, args.key_json, args.docx_out)


if __name__ == "__main__":
    main()
